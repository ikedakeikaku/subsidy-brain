"""Get a usable .docx template for a SubsidyProfile.

Three strategies, tried in order:

  1. **Use cached official forms** — if ``GuidelineFetcher`` previously
     downloaded the publishing body's actual 様式.docx, point at that.
  2. **Look in templates/<program_id>/** — supports the optional hand-
     curated case where the user committed an official file locally.
  3. **Synthesise from profile** — build a docx skeleton with
     ``{{section_id}}`` placeholders matching every section the profile
     declares. This is always available, never depends on the network.

Returning the path to a real .docx means the downstream
``tools.template_filler.fill_template`` (placeholder substitution) and
``tools.document_assembler.assemble_document`` can use the same logic
regardless of where the template came from.
"""
from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from schemas.subsidy_profile import SubsidyProfile

logger = logging.getLogger(__name__)


def _is_valid_docx(path: Path) -> bool:
    """Return True iff ``path`` is a real .docx package (a ZIP whose
    contents include word/document.xml). HTML or PDF served behind a
    .docx filename will fail this check."""
    import zipfile

    try:
        with zipfile.ZipFile(path) as z:
            return "word/document.xml" in z.namelist()
    except (zipfile.BadZipFile, OSError):
        return False


def _slug(name: str) -> str:
    import re
    return re.sub(r"\W+", "_", name.lower()).strip("_") or "subsidy"


def _candidate_local_dirs(
    templates_root: Path, program_id: str, canonical_name: str
) -> list[Path]:
    """Return subdirs of ``templates_root`` that plausibly belong to this
    subsidy, in priority order.

    Match strategy:
      1. Exact slug match (case- and separator-insensitive)
      2. Token overlap: subdir shares any normalised token ≥3 chars with
         the program_id or canonical_name (catches "monodukuri-23" vs
         "monodukuri_23" vs "monozukuri_v23").
    """
    if not templates_root.exists():
        return []

    pid_norm = _slug(program_id)
    keywords = {
        t for t in (_slug(program_id) + "_" + _slug(canonical_name)).split("_")
        if len(t) >= 3
    }

    hits: list[Path] = []
    for d in sorted(templates_root.iterdir()):
        if not d.is_dir() or d.name.startswith("."):
            continue
        if _slug(d.name) == pid_norm:
            hits.append(d)
    for d in sorted(templates_root.iterdir()):
        if not d.is_dir() or d.name.startswith(".") or d in hits:
            continue
        tokens = {t for t in _slug(d.name).split("_") if len(t) >= 3}
        if tokens & keywords:
            hits.append(d)
    return hits


class TemplateSynthesizer:
    """Materialise a .docx template for a given profile."""

    agent_id = "#template"
    agent_name = "TemplateSynthesizer"

    def get_template(
        self,
        profile: SubsidyProfile,
        *,
        fetched_form_paths: dict[str, str] | None = None,
        templates_root: Path | str = "templates",
        drafts_root: Path | str = ".cache/drafts",
    ) -> tuple[Path, str]:
        """Return ``(template_path, source)`` where ``source`` is one of:

          * ``"official"`` — file came from GuidelineFetcher / official URL
          * ``"local"`` — file was committed under ``templates/<program_id>/``
            (this is the only thing that ever lives in ``templates/``)
          * ``"draft"`` — a clearly-labelled DRAFT skeleton synthesised from
            the profile. **Always written under ``.cache/drafts/`` so it
            can't be confused with a user-committed official template.**
        """
        # 1. Use the official 様式 if it actually downloaded AND is a valid
        #    .docx package. The filename can lie — publishers sometimes serve
        #    an HTML landing page or a PDF behind a "様式2.docx" URL, and
        #    python-docx will then raise PackageNotFoundError downstream.
        if fetched_form_paths:
            for form_id, path in fetched_form_paths.items():
                p = Path(path)
                if not (p.exists() and p.stat().st_size > 0):
                    continue
                if p.suffix.lower() != ".docx":
                    continue
                if not _is_valid_docx(p):
                    logger.warning(
                        "TemplateSynthesizer: %s is not a real .docx package "
                        "(probably an HTML page or PDF served at the URL); "
                        "skipping",
                        p,
                    )
                    continue
                if "様式2" in form_id or "経営計画" in form_id or "事業計画" in form_id:
                    logger.info(
                        "TemplateSynthesizer: using official form %s -> %s",
                        form_id,
                        p,
                    )
                    return p, "official"

        # 2. Look in templates/<program_id>/ — only user-committed official
        #    forms live here. Runtime synthesised drafts never write here.
        # The discoverer is non-deterministic about separator style
        # ("monodukuri-23" / "monodukuri_23" can both surface for the same
        # subsidy), so we accept any subdir whose normalised tokens
        # intersect with the program_id's.
        for cand_dir in _candidate_local_dirs(
            Path(templates_root), profile.program_id, profile.canonical_name
        ):
            for cand in sorted(cand_dir.glob("様式*.docx")):
                if _is_valid_docx(cand):
                    return cand, "local"
            for cand in sorted(cand_dir.glob("*.docx")):
                if cand.name.startswith("_"):
                    continue
                if _is_valid_docx(cand):
                    return cand, "local"

        # 3. Synthesise a DRAFT skeleton under .cache/drafts/<program_id>/
        #    The warning header inside the docx and the explicit "draft"
        #    source make it impossible to mistake for the official 様式.
        out = Path(drafts_root) / profile.program_id / "_draft_skeleton.docx"
        return self._synthesise(profile, out), "draft"

    # ------------------------------------------------------------------

    @staticmethod
    def _synthesise(profile: SubsidyProfile, out_path: Path) -> Path:
        out_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Hiragino Sans"
        style.font.size = Pt(10.5)

        section = doc.sections[0]
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        # Warning header — never let this skeleton be confused for the
        # publishing body's official 様式.
        warn = doc.add_paragraph()
        warn_run = warn.add_run(
            "【DRAFT — 本ファイルは公式 様式 ではありません】"
        )
        warn_run.bold = True
        doc.add_paragraph(
            "subsidy-brain が公募要領を読み取って合成した草稿スケルトンです。"
            "実申請時は publishing body（中小企業庁・全国商工会連合会 等）が"
            "配布する公式 様式 docx を取得し、本文だけを移植してください。"
            "本ファイルの罫線・ヘッダ・ページ設定は公式 様式 とは一致しません。"
        )

        title = doc.add_heading(
            f"{profile.canonical_name} 申請書（DRAFT SKELETON）", level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Applicant header
        doc.add_heading("申請者情報", level=1)
        hdr = doc.add_table(rows=4, cols=2)
        hdr.style = "Light Grid Accent 1"
        hdr.rows[0].cells[0].text = "事業者名"
        hdr.rows[0].cells[1].text = "{{ applicant_name }}"
        hdr.rows[1].cells[0].text = "代表者氏名"
        hdr.rows[1].cells[1].text = "{{ representative }}"
        hdr.rows[2].cells[0].text = "事業実施場所"
        hdr.rows[2].cells[1].text = "{{ business_address }}"
        hdr.rows[3].cells[0].text = "従業員数"
        hdr.rows[3].cells[1].text = "{{ employee_count }}"

        # One Heading + body per profile section
        for spec in profile.sections:
            doc.add_heading(spec.display_name, level=1)
            doc.add_paragraph(f"{{{{ {spec.section_id} }}}}")

        doc.save(str(out_path))
        return out_path
