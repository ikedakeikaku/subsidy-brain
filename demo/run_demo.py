"""End-to-end demo: synthetic SMB profile -> subsidy application draft (.docx).

Run modes:
  - `--live`  : actually call Claude API (needs ANTHROPIC_API_KEY)
  - default   : offline mock; uses pre-baked story text. Lets CI run with no key.

Output:
  demo/output/sample_application.docx
"""
from __future__ import annotations

import argparse
import asyncio
import json
import os
import sys
from datetime import datetime
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


# ----------------------------------------------------------------------------
# Story building — either via Claude or via canned mock
# ----------------------------------------------------------------------------

MOCK_STORY = {
    "company_overview": (
        "株式会社サンプル珈琲は、東京都杉並区の住宅街で2018年に開業した"
        "スペシャルティコーヒー専門店です。自家焙煎の品質と店主が産地で"
        "直接買い付けるルートを強みに、平日は近隣住民、週末は遠方の珈琲"
        "愛好家を集めています。席数12席の小規模店舗でありながら、平日"
        "午前のテイクアウトリピート率は60%を超え、固定客に支えられています。"
    ),
    "sales_situation": (
        "売上は2022年18,500千円→2023年21,300千円→2024年22,800千円と"
        "3期連続で成長しているが、店舗キャパ（席数12席）に天井があり、"
        "売上拡大は限界に近づいている。一方、物価高騰で生豆原価が前年比"
        "+18%、現状価格転嫁できていないため営業利益率は伸び悩んでいる。"
    ),
    "challenges": (
        "課題①: 店舗キャパ上限による売上の頭打ち。"
        "課題②: 自家焙煎豆の通販ECがなく遠方需要を逃している。"
        "課題③: 新規顧客のリピート転換率が30%と低い。"
        "課題④: 物価高騰の影響を価格転嫁できていない。"
    ),
    "strategy": (
        "上記課題を解決するため、自家焙煎豆の通販ECサイトを立ち上げ、"
        "店舗売上に依存しない販路を確立する。"
        "施策①: ECサイト構築（Shopify）と商品撮影。"
        "施策②: ブランド統一のためのパッケージデザイン刷新。"
        "施策③: 既存来店客向け試飲会・ワークショップでファン化。"
        "施策④: SNS広告で関東圏外の珈琲愛好家にリーチ。"
    ),
    "expected_outcome": (
        "既存来店客の半数（約720名）を通販リピーター化する。"
        "通販売上で年商+25%、2025年度22,800千円→2026年度28,500千円を目指す。"
        "法人ギフト需要（中元・歳暮）で3社獲得、単価10万円×3＝30万円を確保。"
    ),
    "bonus_env_change": (
        "ロシア・ウクライナ情勢を背景とした輸送コスト上昇、円安、産地国の"
        "気候変動による生豆価格の高騰により、当社の主原料である生豆の仕入"
        "原価は前年比+18%まで上昇している。住宅街立地という性質上、顧客の"
        "多くが価格に敏感な日常使いの利用者であり、価格転嫁による客離れの"
        "リスクが高く、現状価格を据え置かざるを得ない。営業利益率は2024年"
        "6.4%→2025年見込4.8%まで悪化している。"
    ),
}


async def build_story_live(company: dict, guideline_text: str) -> dict:
    """Call Claude through the project's structured-output helper."""
    from tools.claude_client import call_claude_json

    system_prompt = (
        "あなたは小規模事業者の補助金申請書を書く専門コンサルタントです。"
        "公募要領の不変条件（経費区分・文字数上限・審査基準）を必ず守り、"
        "事業者固有のデータ（自社強み・課題・経費計画）に基づいたストーリーを構築してください。"
        "出力は必ず指定のJSON Schemaに従ってください。"
    )

    schema = {
        "type": "object",
        "properties": {
            "company_overview": {"type": "string", "description": "1-1. 自社の概要（600字以内）"},
            "sales_situation": {"type": "string", "description": "1-2. 売上・利益の状況（800字以内）"},
            "challenges": {"type": "string", "description": "1-3. 経営課題（600字以内）"},
            "strategy": {"type": "string", "description": "4-2. 今後のプラン（1000字以内）"},
            "expected_outcome": {"type": "string", "description": "補助事業の効果（500字以内）"},
            "bonus_env_change": {"type": "string", "description": "事業環境変化加点本文（500字以内）"},
        },
        "required": [
            "company_overview", "sales_situation", "challenges",
            "strategy", "expected_outcome", "bonus_env_change",
        ],
    }

    user_message = (
        f"## 公募要領\n{guideline_text}\n\n"
        f"## 申請事業者プロファイル\n{yaml.dump(company, allow_unicode=True)}\n\n"
        "上記の事業者プロファイルと公募要領に基づいて、申請書の主要セクションを書いてください。"
    )

    return await call_claude_json(
        system_prompt=system_prompt,
        user_message=user_message,
        json_schema=schema,
        tool_name="build_application_story",
    )


# ----------------------------------------------------------------------------
# Word output
# ----------------------------------------------------------------------------

def write_word_doc(company: dict, story: dict, out_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Title
    title = doc.add_heading("販路開拓支援補助金 申請書（デモ）", level=0)
    title.alignment = 1

    # Generated metadata
    meta = doc.add_paragraph()
    meta.add_run(f"事業者：{company['company']['name']}\n").bold = True
    meta.add_run(f"代表者：{company['company']['representative']}\n")
    meta.add_run(f"生成日時：{datetime.now().isoformat(timespec='seconds')}\n")
    meta.add_run(
        "※ 本ドキュメントは subsidy-brain のデモ出力です。事業者情報は架空のものです。\n"
    )

    # ＜加点項目＞
    doc.add_heading("＜加点項目＞", level=1)
    doc.add_heading("【重点政策加点】事業環境変化加点", level=2)
    doc.add_paragraph(story["bonus_env_change"])

    # ＜経営計画＞
    doc.add_heading("＜経営計画＞", level=1)
    doc.add_heading("1-1. 自社の概要", level=2)
    doc.add_paragraph(story["company_overview"])

    doc.add_heading("1-2. 売上・利益の状況", level=2)
    doc.add_paragraph(story["sales_situation"])
    # Sales table
    pl = company.get("financial", {}).get("past_3y_pl", [])
    if pl:
        table = doc.add_table(rows=1 + len(pl), cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "年度"
        hdr[1].text = "売上（円）"
        hdr[2].text = "営業利益（円）"
        for i, row in enumerate(pl, start=1):
            cells = table.rows[i].cells
            cells[0].text = str(row["year"])
            cells[1].text = f"{row['revenue']:,}"
            cells[2].text = f"{row['operating_profit']:,}"

    doc.add_heading("1-3. 経営課題", level=2)
    doc.add_paragraph(story["challenges"])

    # ＜補助事業計画＞
    doc.add_heading("＜補助事業計画＞", level=1)
    doc.add_heading("4-2. 今後のプラン", level=2)
    doc.add_paragraph(story["strategy"])

    doc.add_heading("補助事業の効果", level=2)
    doc.add_paragraph(story["expected_outcome"])

    # 経費明細
    doc.add_heading("＜経費明細＞", level=1)
    expenses = company.get("expenses", {})
    breakdown = expenses.get("breakdown", [])
    table = doc.add_table(rows=1 + len(breakdown) + 1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "区分"
    hdr[1].text = "明細"
    hdr[2].text = "金額（円）"
    for i, item in enumerate(breakdown, start=1):
        cells = table.rows[i].cells
        cells[0].text = item["category"]
        cells[1].text = item["item"]
        cells[2].text = f"{item['amount']:,}"
    total_cells = table.rows[-1].cells
    total_cells[0].text = "合計"
    total_cells[1].text = ""
    total_cells[2].text = f"{expenses.get('total', 0):,}"

    p = doc.add_paragraph()
    p.add_run(f"補助金申請額: {expenses.get('subsidy_amount', 0):,}円\n").bold = True
    p.add_run(f"自己負担額  : {expenses.get('self_funding', 0):,}円")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# ----------------------------------------------------------------------------
# Entrypoint
# ----------------------------------------------------------------------------

async def main_async(live: bool) -> None:
    company = yaml.safe_load((ROOT / "demo" / "sample_company.yaml").read_text(encoding="utf-8"))
    guideline = (ROOT / "demo" / "sample_guideline.md").read_text(encoding="utf-8")

    if live:
        if not os.getenv("ANTHROPIC_API_KEY"):
            print("ERROR: --live requires ANTHROPIC_API_KEY env var.", file=sys.stderr)
            sys.exit(2)
        print("→ Calling Claude API for story generation...")
        story = await build_story_live(company, guideline)
    else:
        print("→ Offline mock mode (pass --live to call Claude).")
        story = MOCK_STORY

    out = ROOT / "demo" / "output" / "sample_application.docx"
    write_word_doc(company, story, out)
    print(f"✓ Generated: {out.relative_to(ROOT)}")
    print(f"  Story sections: {list(story.keys())}")

    # Also write the story JSON next to the docx for inspection
    (out.with_suffix(".story.json")).write_text(
        json.dumps(story, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--live",
        action="store_true",
        help="Call Claude API (needs ANTHROPIC_API_KEY); else use canned mock.",
    )
    args = parser.parse_args()
    asyncio.run(main_async(args.live))


if __name__ == "__main__":
    main()
