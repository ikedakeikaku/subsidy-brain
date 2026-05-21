"""Placeholder-substitution Word filler.

The official 様式 docx files distributed by Japanese government bodies have
intricate table layouts, page settings, and built-in styles. Building those
from scratch in python-docx is fragile and almost always loses fidelity.
Instead, we treat the official .docx as a template and replace
``{{placeholder}}`` tokens in place. Tables, borders, fonts and page setup
are preserved because we never reconstruct them.

Usage::

    fill_template(
        template_path="templates/sample_hanro_kaitaku_v1/様式2.docx",
        out_path="demo/output/filled.docx",
        substitutions={
            "section_1_1": "...",
            "section_4_2": "...",
        },
    )

A placeholder is the literal string ``{{ KEY }}`` (spaces optional). Missing
placeholders are logged and replaced with empty string so the rendering
doesn't fail on partially-populated drafts.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_\-\.]+)\s*\}\}")


def fill_template(
    template_path: str | Path,
    out_path: str | Path,
    substitutions: dict[str, str],
) -> dict[str, int]:
    """Open the template, replace placeholders in place, and save.

    Returns a report dict::

        {
          "replaced": <total replacement count>,
          "unique_keys_used": <number of distinct keys substituted>,
          "missing_keys": [<placeholders in doc with no substitution>],
        }
    """
    template_path = Path(template_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template_path))
    counts: dict[str, int] = {}

    # Paragraphs (top-level body)
    for para in doc.paragraphs:
        counts = _substitute_in_paragraph(para, substitutions, counts)

    # Tables — note: a table cell can contain paragraphs (recursively for
    # nested tables), so we recurse via _substitute_in_cell.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                counts = _substitute_in_cell(cell, substitutions, counts)

    # Headers & footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            for para in header_footer.paragraphs:
                counts = _substitute_in_paragraph(para, substitutions, counts)

    doc.save(str(out_path))

    missing = _scan_unresolved_placeholders(out_path)

    return {
        "replaced": sum(counts.values()),
        "unique_keys_used": len(counts),
        "missing_keys": sorted(missing),
    }


def _substitute_in_paragraph(
    paragraph, substitutions: dict[str, str], counts: dict[str, int]
) -> dict[str, int]:
    if not paragraph.text or "{{" not in paragraph.text:
        return counts

    full_text = paragraph.text
    new_text, replaced = _apply(full_text, substitutions)
    if not replaced:
        return counts

    # python-docx splits text into runs (formatting boundaries). The simplest
    # robust replacement collapses runs into one whose text is the replaced
    # string. This loses per-run formatting *within* the replaced region but
    # preserves the surrounding paragraph layout (alignment, indent, list
    # marker, etc.).
    for run in list(paragraph.runs):
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)

    for key, _val in replaced:
        counts[key] = counts.get(key, 0) + 1
    return counts


def _substitute_in_cell(
    cell, substitutions: dict[str, str], counts: dict[str, int]
) -> dict[str, int]:
    for para in cell.paragraphs:
        counts = _substitute_in_paragraph(para, substitutions, counts)
    for tbl in cell.tables:
        for row in tbl.rows:
            for inner in row.cells:
                counts = _substitute_in_cell(inner, substitutions, counts)
    return counts


def _apply(
    text: str, substitutions: dict[str, str]
) -> tuple[str, list[tuple[str, str]]]:
    replaced: list[tuple[str, str]] = []

    def _replace(match: re.Match) -> str:
        key = match.group(1)
        if key in substitutions:
            replaced.append((key, substitutions[key]))
            return substitutions[key]
        logger.debug("placeholder %s has no substitution; left empty", key)
        replaced.append((key, ""))
        return ""

    new_text = _PLACEHOLDER_RE.sub(_replace, text)
    return new_text, replaced


def _scan_unresolved_placeholders(docx_path: Path) -> set[str]:
    doc = Document(str(docx_path))
    leftovers: set[str] = set()
    for para in doc.paragraphs:
        leftovers.update(_PLACEHOLDER_RE.findall(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    leftovers.update(_PLACEHOLDER_RE.findall(para.text))
    return leftovers
