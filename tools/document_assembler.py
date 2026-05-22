"""Profile-driven document assembler.

Builds the final application .docx by:

  1. Generating every chart declared in the profile (matplotlib → PNG).
  2. Generating every table declared in the profile (Pydantic data → rows).
  3. Building the document from scratch following the profile's section
     order, then inserting charts / tables at the declared placement.

This replaces the simple ``fill_template`` approach when a profile is given
because we have full structural information and don't need a pre-existing
.docx skeleton. The result is a faithful rendition of the subsidy's
expected structure with real charts and tables instead of placeholders.
"""
from __future__ import annotations

import logging
import tempfile
from pathlib import Path
from typing import Any

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from schemas.subsidy_profile import ChartSpec, ChartType, SubsidyProfile, TableSpec, TableType

matplotlib.use("Agg")
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Chart generation
# ---------------------------------------------------------------------------


def _setup_japanese_font() -> None:
    """Pick a font that renders Japanese without tofu (mojibake squares)."""
    import matplotlib.font_manager as fm

    candidates = [
        "Hiragino Sans",
        "Hiragino Maru Gothic Pro",
        "Yu Gothic",
        "Noto Sans CJK JP",
        "IPAGothic",
    ]
    available = {f.name for f in fm.fontManager.ttflist}
    for name in candidates:
        if name in available:
            plt.rcParams["font.family"] = name
            return


def _generate_chart(
    chart: ChartSpec, data: Any, out_path: Path
) -> Path | None:
    """Render one chart based on its declared type. Returns None if data is empty."""
    _setup_japanese_font()
    fig, ax = plt.subplots(figsize=(7.5, 3.2), dpi=120)

    try:
        if chart.chart_type == ChartType.REVENUE_TREND:
            rows = data or []
            if not rows:
                plt.close(fig)
                return None
            labels = [str(r.get("year", "")) for r in rows]
            revenues = [r.get("revenue", 0) / 10_000 for r in rows]
            bars = ax.bar(labels, revenues, color="#4c78a8", edgecolor="#1f4e79")
            ax.set_ylabel("売上（万円）")
            ax.set_title(chart.title)
            for bar, val in zip(bars, revenues, strict=False):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height(),
                    f"{val:,.0f}",
                    ha="center",
                    va="bottom",
                    fontsize=9,
                )

        elif chart.chart_type == ChartType.EFFECT_BEFORE_AFTER:
            # data should be like {"before": int, "after": int}
            # if not provided, fall back to the company financial trend
            before = (data or {}).get("before")
            after = (data or {}).get("after")
            if before is None or after is None:
                plt.close(fig)
                return None
            labels = ["補助事業実施前", "補助事業実施後"]
            values = [before / 10_000, after / 10_000]
            colors = ["#a8a8a8", "#4c78a8"]
            bars = ax.bar(labels, values, color=colors, edgecolor="#444")
            ax.set_ylabel("年商（万円）")
            ax.set_title(chart.title)
            growth = (after - before) / before * 100 if before else 0.0
            ax.text(
                0.95,
                0.95,
                f"成長率 {growth:+.1f}%",
                transform=ax.transAxes,
                ha="right",
                va="top",
                fontsize=11,
                bbox={"facecolor": "white", "alpha": 0.8, "edgecolor": "none"},
            )
            for bar, val in zip(bars, values, strict=False):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height(),
                    f"{val:,.0f}",
                    ha="center",
                    va="bottom",
                    fontsize=9,
                )

        else:
            logger.warning("chart_type %s not yet implemented", chart.chart_type)
            plt.close(fig)
            return None

        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        fig.tight_layout()
        fig.savefig(str(out_path), bbox_inches="tight", dpi=120)
    finally:
        plt.close(fig)
    return out_path


# ---------------------------------------------------------------------------
# Table generation from profile + data
# ---------------------------------------------------------------------------


def _add_table_to_doc(doc: Document, table_spec: TableSpec, data: Any) -> None:
    rows: list[list[str]] = []

    if table_spec.table_type == TableType.PL_HISTORY:
        for r in data or []:
            rows.append(
                [
                    str(r.get("year", "")),
                    f"{r.get('revenue', 0):,}",
                    f"{r.get('operating_profit', 0):,}",
                ]
            )

    elif table_spec.table_type == TableType.EXPENSE_BREAKDOWN:
        for r in data or []:
            rows.append(
                [
                    str(r.get("category", "")),
                    str(r.get("item", "")),
                    f"{r.get('amount', 0):,}",
                ]
            )

    elif table_spec.table_type == TableType.SCHEDULE:
        for r in data or []:
            rows.append(
                [
                    str(r.get("when", "")),
                    str(r.get("item", "")),
                    str(r.get("content", "")),
                ]
            )

    elif table_spec.table_type == TableType.KPI_TABLE:
        for r in data or []:
            rows.append([str(r.get(c, "")) for c in table_spec.columns])

    if not rows:
        doc.add_paragraph(f"（{table_spec.title}：データなし）")
        return

    p = doc.add_paragraph()
    p.add_run(f"【{table_spec.title}】").bold = True

    word_table = doc.add_table(rows=1 + len(rows), cols=len(table_spec.columns))
    word_table.style = "Light Grid Accent 1"
    for i, col in enumerate(table_spec.columns):
        cell = word_table.rows[0].cells[i]
        cell.text = col
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            word_table.rows[i].cells[j].text = val


# ---------------------------------------------------------------------------
# Main assembler
# ---------------------------------------------------------------------------


def _get(data: Any, dotted: str) -> Any:
    """Resolve 'a.b.c' on nested dicts. Returns None if any key is missing."""
    cur: Any = data
    for key in dotted.split("."):
        if isinstance(cur, dict):
            cur = cur.get(key)
        else:
            return None
        if cur is None:
            return None
    return cur


def assemble_document(
    profile: SubsidyProfile,
    company: dict[str, Any],
    story: dict[str, str],
    out_path: Path,
    *,
    extra_metadata: dict[str, str] | None = None,
    quality_block: str | None = None,
) -> dict[str, Any]:
    """Build the final application .docx.

    Returns a report containing the list of inserted charts / tables and the
    list of sections rendered.
    """
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Title block
    title = doc.add_heading(profile.canonical_name + " 申請書（自動生成）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    if extra_metadata:
        for k, v in extra_metadata.items():
            meta.add_run(f"{k}: {v}\n")

    # Applicant header
    company_block = company.get("company", {})
    doc.add_heading("申請者情報", level=1)
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Light Grid Accent 1"
    hdr.rows[0].cells[0].text = "事業者名"
    hdr.rows[0].cells[1].text = str(company_block.get("name", ""))
    hdr.rows[1].cells[0].text = "代表者氏名"
    hdr.rows[1].cells[1].text = str(company_block.get("representative", ""))
    hdr.rows[2].cells[0].text = "事業実施場所"
    hdr.rows[2].cells[1].text = (
        f"{company_block.get('prefecture', '')}{company_block.get('city', '')}"
    )
    hdr.rows[3].cells[0].text = "従業員数"
    hdr.rows[3].cells[1].text = f"{company_block.get('employees', '')}名"

    # ＜加点項目＞ block — render before the main sections (matches the
    # layout of typical 様式2 forms). bonus_items contents come from the
    # BonusEvaluator and are addressed in story under "bonus_{item_id}".
    bonus_rendered: list[str] = []
    if profile.bonus_items:
        doc.add_heading("＜加点項目＞", level=1)
        for item in profile.bonus_items:
            key = f"bonus_{item.item_id}"
            body = (story or {}).get(key, "")
            if not body:
                continue  # not applicable for this company
            label = f"【{item.category}】{item.display_name}" if item.category else item.display_name
            doc.add_heading(label, level=2)
            doc.add_paragraph(body)
            bonus_rendered.append(item.item_id)

    inserted_charts: list[str] = []
    inserted_tables: list[str] = []

    with tempfile.TemporaryDirectory() as tmp:
        tmp_root = Path(tmp)

        for section_spec in profile.sections:
            doc.add_heading(section_spec.display_name, level=1)
            text = story.get(section_spec.section_id, "")
            doc.add_paragraph(text)

            # Insert any chart that requests placement after this section
            for chart in profile.charts:
                if chart.place_after_section != section_spec.section_id:
                    continue
                data = _get(company, chart.data_path)
                if chart.chart_type == ChartType.EFFECT_BEFORE_AFTER:
                    # Auto-build before/after from the financial trend
                    pl = _get(company, "financial.past_3y_pl") or []
                    expected = (
                        _get(company, "planned_project.expected_outcomes") or []
                    )
                    if pl:
                        last = pl[-1].get("revenue", 0)
                        # Look for an explicit target in expected_outcomes,
                        # else fall back to +25%
                        after = int(last * 1.25)
                        for item in expected:
                            if isinstance(item, dict) and "target_revenue" in item:
                                after = int(item["target_revenue"])
                                break
                        data = {"before": last, "after": after}
                out_png = tmp_root / f"{chart.chart_id}.png"
                rendered = _generate_chart(chart, data, out_png)
                if rendered:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(str(rendered), width=Cm(chart.width_cm))
                    cap = doc.add_paragraph(f"図. {chart.title}")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    inserted_charts.append(chart.chart_id)

            # Insert any table that requests placement after this section
            for table_spec in profile.tables:
                if table_spec.place_after_section != section_spec.section_id:
                    continue
                data = _get(company, table_spec.data_path)
                _add_table_to_doc(doc, table_spec, data)
                inserted_tables.append(table_spec.table_id)

        # Optional: quality score block at the end
        if quality_block:
            doc.add_heading("自己採点（subsidy-brain による品質チェック結果）", level=1)
            doc.add_paragraph(quality_block)

        doc.save(str(out_path))

    return {
        "output": str(out_path),
        "sections_rendered": [s.section_id for s in profile.sections],
        "charts_inserted": inserted_charts,
        "tables_inserted": inserted_tables,
        "bonus_items_rendered": bonus_rendered,
    }
