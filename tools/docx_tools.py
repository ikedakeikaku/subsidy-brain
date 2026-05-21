"""Word文書（.docx）生成ユーティリティ

python-docx を使用したテンプレート読み込み・プレースホルダー置換・
テーブル挿入・画像挿入・保存を提供する。
"""

import logging
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

logger = logging.getLogger(__name__)

# プロジェクトルートの templates/ ディレクトリ
_TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"

# テンプレートIDからファイル名へのマッピング
_TEMPLATE_MAP: dict[str, str] = {
    "application_form": "application_form.docx",
    "expense_table": "expense_table.docx",
}


def load_template(template_id: str) -> Document:
    """テンプレートIDに対応するWordテンプレートを読み込む。

    Args:
        template_id: テンプレート識別子（例: "application_form"）。

    Returns:
        読み込まれた Document オブジェクト。

    Raises:
        FileNotFoundError: テンプレートファイルが存在しない場合。
        ValueError: 未知のテンプレートIDの場合。
    """
    filename = _TEMPLATE_MAP.get(template_id)
    if filename is None:
        # マップにない場合はそのまま .docx を付けて探す
        filename = f"{template_id}.docx"

    template_path = _TEMPLATES_DIR / filename
    if not template_path.exists():
        raise FileNotFoundError(
            f"テンプレートが見つかりません: {template_path}"
        )

    logger.info("テンプレート読み込み: %s", template_path)
    return Document(str(template_path))


def _replace_in_paragraph(paragraph, placeholder: str, text: str) -> bool:
    """段落内のプレースホルダーを置換する。書式を保持する。

    プレースホルダーが複数のrunに分割されている場合にも対応する。

    Returns:
        置換が行われた場合 True。
    """
    full_text = paragraph.text
    if placeholder not in full_text:
        return False

    # プレースホルダーが単一のrunに収まっている場合
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, text)
            return True

    # 複数のrunに分割されている場合: 全runのテキストを結合し再構築
    # 最初のrunに置換後テキストを設定し、残りのrunをクリア
    new_text = full_text.replace(placeholder, text)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def fill_section(
    doc: Document,
    placeholder: str,
    text: str,
    max_chars: int | None = None,
) -> int:
    """ドキュメント内のプレースホルダーをテキストで置換する。

    段落およびテーブルセル内のプレースホルダー ``{{placeholder}}`` を検索し、
    指定テキストで置換する。元の段落の書式（フォント・スタイル）を保持する。

    Args:
        doc: 対象の Document オブジェクト。
        placeholder: 置換対象のプレースホルダー名（二重波括弧なし）。
            例: "company_name" → ドキュメント内の ``{{company_name}}`` を置換。
        text: 挿入するテキスト。
        max_chars: テキストの最大文字数。超過分は切り詰められる。

    Returns:
        挿入されたテキストの文字数。
    """
    tag = "{{" + placeholder + "}}"

    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars]
        logger.warning(
            "テキストを %d 文字に切り詰めました (placeholder=%s)", max_chars, placeholder
        )

    replaced = False

    # 本文の段落を検索
    for paragraph in doc.paragraphs:
        if _replace_in_paragraph(paragraph, tag, text):
            replaced = True

    # テーブルセル内の段落も検索
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if _replace_in_paragraph(paragraph, tag, text):
                        replaced = True

    if not replaced:
        logger.warning("プレースホルダー '%s' が見つかりませんでした", tag)

    return len(text)


def insert_table(
    doc: Document,
    placeholder: str,
    headers: list[str],
    rows: list[list[str]],
) -> None:
    """プレースホルダーの位置にテーブルを挿入する。

    プレースホルダーを含む段落を見つけ、その直後にテーブルを追加し、
    元の段落（プレースホルダーテキスト）を削除する。

    Args:
        doc: 対象の Document オブジェクト。
        placeholder: テーブル挿入位置のプレースホルダー名。
        headers: テーブルのヘッダー行（列名リスト）。
        rows: テーブルのデータ行リスト。
    """
    tag = "{{" + placeholder + "}}"

    for i, paragraph in enumerate(doc.paragraphs):
        if tag not in paragraph.text:
            continue

        # テーブルを作成（ヘッダー + データ行）
        num_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = "Table Grid"

        # ヘッダー行
        for col_idx, header in enumerate(headers):
            table.rows[0].cells[col_idx].text = header

        # データ行
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    table.rows[row_idx].cells[col_idx].text = str(cell_text)

        # テーブル要素をプレースホルダー段落の直後に移動
        paragraph._element.addnext(table._tbl)

        # プレースホルダー段落のテキストをクリア
        for run in paragraph.runs:
            run.text = ""
        paragraph.text = ""

        logger.info(
            "テーブル挿入完了 (placeholder=%s): %d 列 x %d 行",
            placeholder,
            num_cols,
            len(rows),
        )
        return

    logger.warning("テーブル挿入用プレースホルダー '%s' が見つかりませんでした", tag)


def insert_image(
    doc: Document,
    placeholder: str,
    image_path: str,
    width_cm: float = 15.0,
) -> None:
    """プレースホルダーの位置に画像を挿入する。

    Args:
        doc: 対象の Document オブジェクト。
        placeholder: 画像挿入位置のプレースホルダー名。
        image_path: 挿入する画像ファイルのパス。
        width_cm: 画像の表示幅（cm）。

    Raises:
        FileNotFoundError: 画像ファイルが存在しない場合。
    """
    tag = "{{" + placeholder + "}}"

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"画像ファイルが見つかりません: {image_path}")

    for paragraph in doc.paragraphs:
        if tag not in paragraph.text:
            continue

        # プレースホルダーテキストをクリア
        for run in paragraph.runs:
            run.text = ""
        paragraph.text = ""

        # 画像を挿入
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(width_cm))

        logger.info("画像挿入完了 (placeholder=%s): %s", placeholder, image_path)
        return

    logger.warning("画像挿入用プレースホルダー '%s' が見つかりませんでした", tag)


def save_document(doc: Document, output_path: str) -> str:
    """Documentを指定パスに保存する。

    出力ディレクトリが存在しない場合は自動作成する。

    Args:
        doc: 保存する Document オブジェクト。
        output_path: 出力ファイルパス。

    Returns:
        保存先の絶対パス。
    """
    output = Path(output_path).resolve()
    output.parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output))
    logger.info("ドキュメント保存完了: %s", output)
    return str(output)


def fill_section_formatted(
    doc: Document,
    placeholder: str,
    text: str,
    bold_terms: list[str] | None = None,
    underline_terms: list[str] | None = None,
    max_chars: int | None = None,
) -> int:
    """プレースホルダーを書式付きテキストで置換する。

    指定したキーワードを太字・アンダーライン付きで挿入できる。
    bold_terms と underline_terms の両方に含まれるタームは両方が適用される。

    Args:
        doc: 対象の Document オブジェクト。
        placeholder: 置換対象のプレースホルダー名（二重波括弧なし）。
        text: 挿入するテキスト。
        bold_terms: 太字にするキーワードのリスト。
        underline_terms: アンダーラインを付けるキーワードのリスト。
        max_chars: テキストの最大文字数。超過分は切り詰められる。

    Returns:
        挿入されたテキストの文字数。
    """
    tag = "{{" + placeholder + "}}"
    bold_terms = bold_terms or []
    underline_terms = underline_terms or []

    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars]
        logger.warning(
            "テキストを %d 文字に切り詰めました (placeholder=%s)", max_chars, placeholder
        )

    replaced = False

    for paragraph in doc.paragraphs:
        if tag not in paragraph.text:
            continue

        # 既存 run の書式を参照し段落をクリア
        base_font_name = None
        base_font_size = None
        if paragraph.runs:
            first_run = paragraph.runs[0]
            base_font_name = first_run.font.name
            base_font_size = first_run.font.size

        for run in paragraph.runs:
            run.text = ""

        if not bold_terms and not underline_terms:
            # 書式指定なし: シンプルに置換
            run = paragraph.add_run(text)
            if base_font_name:
                run.font.name = base_font_name
            if base_font_size:
                run.font.size = base_font_size
        else:
            # テキストをキーワードで分割し run ごとに書式を適用
            all_terms = list(set(bold_terms + underline_terms))
            # 最長一致優先でトークン分割
            import re
            pattern = "(" + "|".join(re.escape(t) for t in sorted(all_terms, key=len, reverse=True)) + ")"
            parts = re.split(pattern, text)
            for part in parts:
                if not part:
                    continue
                run = paragraph.add_run(part)
                if base_font_name:
                    run.font.name = base_font_name
                if base_font_size:
                    run.font.size = base_font_size
                if part in bold_terms:
                    run.bold = True
                if part in underline_terms:
                    run.underline = True

        replaced = True
        logger.info("書式付き置換完了 (placeholder=%s)", placeholder)

    if not replaced:
        logger.warning("プレースホルダー '%s' が見つかりませんでした", tag)

    return len(text)


def _set_cell_shading(cell, fill_rgb: tuple[int, int, int]) -> None:
    """テーブルセルに背景色を設定する (OxmlElement 経由)。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = "{:02X}{:02X}{:02X}".format(*fill_rgb)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def insert_styled_table(
    doc: Document,
    placeholder: str,
    headers: list[str],
    rows: list[list[str]],
    header_bg_rgb: tuple[int, int, int] = (217, 225, 242),
) -> None:
    """スタイル付きテーブルをプレースホルダー位置に挿入する。

    ヘッダー行に背景色・太字・白文字を適用し、全セルに罫線を付ける。
    数値（半角数字のみ）のセルは右寄せにする。

    Args:
        doc: 対象の Document オブジェクト。
        placeholder: テーブル挿入位置のプレースホルダー名。
        headers: テーブルのヘッダー行（列名リスト）。
        rows: テーブルのデータ行リスト。
        header_bg_rgb: ヘッダー背景色 (R, G, B)。デフォルトは淡い青。
    """
    tag = "{{" + placeholder + "}}"

    for paragraph in doc.paragraphs:
        if tag not in paragraph.text:
            continue

        num_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.style = "Table Grid"

        # ヘッダー行の書式設定
        header_row = table.rows[0]
        for col_idx, header in enumerate(headers):
            cell = header_row.cells[col_idx]
            cell.text = ""
            _set_cell_shading(cell, header_bg_rgb)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

        # データ行
        for row_idx, row_data in enumerate(rows, start=1):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= num_cols:
                    break
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_text)
                # 数値列は右寄せ
                val = str(cell_text).replace(",", "").replace("円", "").strip()
                if val.lstrip("-").isdigit():
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # テーブルをプレースホルダー段落の直後に移動
        paragraph._element.addnext(table._tbl)

        # プレースホルダー段落をクリア
        for run in paragraph.runs:
            run.text = ""
        paragraph.text = ""

        logger.info(
            "スタイル付きテーブル挿入完了 (placeholder=%s): %d 列 x %d 行",
            placeholder,
            num_cols,
            len(rows),
        )
        return

    logger.warning("テーブル挿入用プレースホルダー '%s' が見つかりませんでした", tag)


def set_document_style(doc: Document) -> None:
    """ドキュメント全体のスタイルを設定する。

    A4 用紙、余白 20mm、デフォルトフォント MS明朝 10.5pt を適用する。

    Args:
        doc: 対象の Document オブジェクト。
    """
    section = doc.sections[0]

    # A4 サイズ
    section.page_width = Mm(210)
    section.page_height = Mm(297)

    # 余白 20mm
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    # デフォルトフォント設定
    style = doc.styles["Normal"]
    font = style.font
    font.name = "MS明朝"
    font.size = Pt(10.5)

    # CJK フォント設定 (XML レベル)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "MS明朝")

    logger.info("ドキュメントスタイル設定完了 (A4, 余白20mm, MS明朝 10.5pt)")


def add_heading_paragraph(doc: Document, text: str, level: int = 1) -> None:
    """スタイル付き見出し段落をドキュメントに追加する。

    レベルに応じてフォントサイズ・太字・アンダーラインを設定する。

    Args:
        doc: 対象の Document オブジェクト。
        text: 見出しテキスト。
        level: 見出しレベル（1〜3）。
            - 1: 14pt 太字 アンダーライン
            - 2: 12pt 太字
            - 3: 10.5pt 太字
    """
    heading_styles = {
        1: {"size": Pt(14), "bold": True, "underline": True},
        2: {"size": Pt(12), "bold": True, "underline": False},
        3: {"size": Pt(10.5), "bold": True, "underline": False},
    }
    style_props = heading_styles.get(level, heading_styles[1])

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = style_props["size"]
    run.bold = style_props["bold"]
    run.underline = style_props["underline"]

    # CJK フォント指定
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "MS明朝")

    logger.debug("見出し追加 (level=%d): %s", level, text)
