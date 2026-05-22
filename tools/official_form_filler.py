"""Fill an actual official subsidy form (.docx) without breaking formatting.

Real 様式 docx files distributed by Japanese government bodies don't have
``{{placeholder}}`` markers. They have:

  * Heading paragraphs ("1-1. 自社の概要") followed by empty body space.
  * Applicant-info tables with label cells ("事業者名", "代表者氏名") and
    adjacent empty value cells.
  * Pre-formatted tables that the applicant fills in by row.

This module fills such templates by **structural matching** rather than
string substitution. It walks the document, recognises headings, and
inserts content into the right place while leaving every other paragraph,
run, table border, page setting, and embedded image untouched.

Three filling strategies, applied in order:

  1. **Heading-aware section fill** — for every section in the profile,
     locate the matching heading and insert the body text directly after
     it (or into the first empty paragraph below).
  2. **Applicant info table fill** — find tables that contain known
     labels (事業者名, 代表者, 住所, 従業員数) and write into the cell to
     their right.
  3. **Generated table/chart insertion** — if the profile declares
     additional tables (schedule, expense breakdown) and the official
     template doesn't already provide them, append them at the section
     boundary the profile requests.
"""
from __future__ import annotations

import logging
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.shared import Cm
from docx.table import _Cell

from schemas.subsidy_profile import SubsidyProfile, TableSpec, TableType

logger = logging.getLogger(__name__)


# Labels we recognise in applicant-info tables. The right-hand cell of any
# cell whose stripped text matches one of these gets the corresponding
# company value.
_APPLICANT_LABEL_MAP: dict[str, str] = {
    "事業者名": "company.name",
    "申請者名": "company.name",
    "代表者氏名": "company.representative",
    "代表者": "company.representative",
    "事業実施場所": "address",
    "事業所所在地": "address",
    "所在地": "address",
    "従業員数": "company.employees",
    "業種": "company.industry",
    "設立": "company.founded",
}


def fill_official_form(
    template_path: str | Path,
    out_path: str | Path,
    *,
    profile: SubsidyProfile,
    story: dict[str, str],
    company: dict[str, Any],
) -> dict:
    """Open ``template_path``, fill it from ``story`` / ``company``, save.

    Returns a report::

        {
          "sections_filled": [<section_id>, ...],
          "sections_not_found": [<display_name>, ...],
          "applicant_cells_filled": int,
          "tables_appended": [<table_id>, ...],
          "charts_appended": [<chart_id>, ...],
        }
    """
    template_path = Path(template_path)
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc: _Document = Document(str(template_path))

    sections_filled: list[str] = []
    sections_not_found: list[str] = []
    for spec in profile.sections:
        # Containers carry no body, skip them.
        if getattr(spec, "kind", "section") == "container":
            continue
        text = story.get(spec.section_id, "")
        if not text:
            continue
        kind = getattr(spec, "kind", "section")
        if _fill_section_after_heading(doc, spec.display_name, text, kind=kind):
            sections_filled.append(spec.section_id)
        else:
            sections_not_found.append(spec.display_name)

    applicant_cells = _fill_applicant_tables(doc, company)
    tables_appended = _append_profile_tables(doc, profile, company)
    charts_appended = _append_profile_charts(doc, profile, company)

    doc.save(str(out_path))

    return {
        "sections_filled": sections_filled,
        "sections_not_found": sections_not_found,
        "applicant_cells_filled": applicant_cells,
        "tables_appended": tables_appended,
        "charts_appended": charts_appended,
        "template_used": str(template_path),
        "output": str(out_path),
    }


# ---------------------------------------------------------------------------
# Heading-aware section fill
# ---------------------------------------------------------------------------


def _normalise(s: str) -> str:
    """Drop whitespace / common punctuation for fuzzy heading match."""
    return re.sub(r"[\s　．\.\(\)（）\-―—_]+", "", s)


_SECTION_NUMBER_RE = re.compile(r"([0-9]+(?:[-\.][0-9]+)*)")
_BONUS_RE = re.compile(r"加点")
_ENV_CHANGE_RE = re.compile(r"事業環境変化|物価高騰|環境変化")
_WAGE_RE = re.compile(r"賃[金上]+|賃金引[上き]")


def _section_signature(text: str) -> tuple[str | None, str]:
    """Extract a stable signature for a heading.

    Returns ``(section_signature, normalised_text)``. The signature is
    derived from the strongest stable identifier we can find:
      * section number like ``"1-1"`` or ``"4-2"``
      * "bonus:env_change" for any 加点 + 事業環境変化 heading
      * "bonus:wage" for 賃金引上げ加点
      * "bonus:other" for any other 加点 heading
    """
    match = _SECTION_NUMBER_RE.search(text)
    if match:
        return (match.group(1), _normalise(text))
    if _BONUS_RE.search(text):
        if _ENV_CHANGE_RE.search(text):
            return ("bonus:env_change", _normalise(text))
        if _WAGE_RE.search(text):
            return ("bonus:wage", _normalise(text))
        return ("bonus:other", _normalise(text))
    if _ENV_CHANGE_RE.search(text):
        return ("bonus:env_change", _normalise(text))
    return (None, _normalise(text))


def _fill_section_after_heading(
    doc: _Document,
    heading_text: str,
    body: str,
    *,
    kind: str = "section",
) -> bool:
    """Find the paragraph whose text matches ``heading_text`` and write
    ``body`` immediately after it.

    For leaves (the publisher's "（XXX字以内）" inputs), we MUST match on
    the full normalised display text — section-number heuristics return
    spurious matches because several leaves share the same numeric prefix
    or none at all. For top-level sections, the numeric prefix is the
    strongest available signal.

    Match precedence:
      Leaves: full normalised text equality, then substring either direction.
      Sections / others: section-number, bonus tag, substring.
    """
    target_num, target_norm = _section_signature(heading_text)
    paras = list(doc.paragraphs)

    if kind == "leaf":
        # Pass A: exact normalised match (covers the bulk of leaves)
        for i, para in enumerate(paras):
            ptext_norm = _normalise(para.text)
            if not ptext_norm:
                continue
            if ptext_norm == target_norm:
                _insert_leaf_body(para, paras, i, body)
                return True
        # Pass B: substring either direction (rare phrasing drift between
        # what the synthesiser stored and the raw docx paragraph text)
        for i, para in enumerate(paras):
            ptext_norm = _normalise(para.text)
            if not ptext_norm:
                continue
            if target_norm in ptext_norm or ptext_norm in target_norm:
                _insert_leaf_body(para, paras, i, body)
                return True
        return False

    # Pass 1: section-number match
    if target_num is not None:
        for i, para in enumerate(paras):
            if not para.text.strip():
                continue
            num, _ = _section_signature(para.text)
            if num == target_num:
                _insert_body_after(para, paras, i, body)
                return True

    # Pass 2: normalised substring either direction
    for i, para in enumerate(paras):
        ptext_norm = _normalise(para.text)
        if not ptext_norm:
            continue
        if (
            ptext_norm == target_norm
            or target_norm in ptext_norm
            or ptext_norm in target_norm
        ):
            _insert_body_after(para, paras, i, body)
            return True
    return False


_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _insert_body_after(heading_para, paras, idx: int, body: str) -> None:
    """If the paragraph right after ``heading_para`` is empty, fill it.
    Otherwise insert a new paragraph using lxml-level placement so the
    heading-body adjacency is preserved.
    """
    # Try filling the next paragraph if it's empty
    if idx + 1 < len(paras):
        next_para = paras[idx + 1]
        if not next_para.text.strip():
            next_para.add_run(body)
            return
    # Else insert a brand new paragraph after the heading. We do this via
    # lxml to ensure ordering, since python-docx's default add_paragraph
    # appends at the end of the body.
    new_p = heading_para._parent.add_paragraph(body)
    heading_element = heading_para._p
    heading_element.addnext(new_p._p)


def _insert_leaf_body(heading_para, paras, idx: int, body: str) -> None:
    """Insert body for a leaf heading.

    Japanese 様式 docs put each leaf's input region inside a table
    immediately after the heading:

      * **rows=1, single cell**: the cell contains the instruction text
        ("市場・顧客動向を始めとした..."). The body is appended INSIDE
        the cell, after the instruction.
      * **rows=2**: row 0 is the instruction; row 1 is empty. The body
        is written into row 1's first cell.
      * **no following table**: fall back to the legacy paragraph-after
        insertion (covers headings whose 様式 happens to use an empty
        paragraph as the input region instead of a table).

    This is what the publisher actually expects: their format is the
    layout the reviewer sees, and putting body content outside the
    designated table cell breaks the visual binding between the field
    label (with its 字数制限 annotation) and the answer.
    """
    next_tbl = _next_sibling_tbl(heading_para._p)
    if next_tbl is not None:
        rows = next_tbl.findall(f"{_W_NS}tr")
        if len(rows) >= 2:
            # Last row is the input slot; first cell holds the answer.
            target_cells = rows[-1].findall(f"{_W_NS}tc")
            if target_cells:
                _write_body_into_cell(target_cells[0], body)
                return
        if len(rows) == 1:
            target_cells = rows[0].findall(f"{_W_NS}tc")
            if target_cells:
                # Single-cell input region: the publisher's instruction
                # text is placeholder content (like form-control gray
                # placeholder text in HTML). Replace it with the body.
                _write_body_into_cell(target_cells[0], body)
                return
    # No table after the heading — fall through to the paragraph path.
    _insert_body_after(heading_para, paras, idx, body)


def _next_sibling_tbl(p_element):
    """Walk forward from a paragraph element to find the next table.
    Skips empty paragraphs (whitespace, page breaks). Returns ``None`` if
    a non-empty paragraph is reached first."""
    sibling = p_element.getnext()
    while sibling is not None:
        tag = sibling.tag.split("}")[-1]
        if tag == "tbl":
            return sibling
        if tag == "p":
            # Skip blank paragraphs; stop if a real heading-or-text shows up.
            text = "".join(
                (t.text or "")
                for t in sibling.iter(f"{_W_NS}t")
            ).strip()
            if text:
                return None
        sibling = sibling.getnext()
    return None


def _write_body_into_cell(tc_element, body: str) -> None:
    """Replace the cell's content with ``body`` (preserving cell style).

    Used for the row-1 input cell of a 2-row instruction/input table —
    that cell was empty, so we don't have to keep any prior text.
    """
    from docx.oxml.ns import qn

    # Remove existing paragraph contents (but keep cell properties tcPr).
    for child in list(tc_element):
        if child.tag == qn("w:p"):
            tc_element.remove(child)
    # Add one paragraph carrying the body.
    new_p = _make_body_paragraph(body)
    tc_element.append(new_p)


def _make_body_paragraph(body: str):
    """Build a fresh ``<w:p>`` element containing ``body``. Plain Normal
    style; no special run formatting — keeps the cell's own paragraph
    style by default."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = body
    r.append(t)
    p.append(r)
    return p


# ---------------------------------------------------------------------------
# Applicant info table fill
# ---------------------------------------------------------------------------


def _fill_applicant_tables(doc: _Document, company: dict[str, Any]) -> int:
    """Walk every table; when a cell's text matches a known label, write
    the matching company value into the cell to its right (or below)."""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for col_idx, cell in enumerate(cells):
                label_norm = _normalise(cell.text)
                for label, path in _APPLICANT_LABEL_MAP.items():
                    if label_norm == _normalise(label):
                        value = _resolve_path(company, path)
                        if value is not None and col_idx + 1 < len(cells):
                            target = cells[col_idx + 1]
                            if not target.text.strip() or "{{" in target.text:
                                _set_cell_text(target, str(value))
                                count += 1
                        break
    return count


def _resolve_path(data: dict, dotted: str) -> Any:
    if dotted == "address":
        company = data.get("company", {}) or {}
        return f"{company.get('prefecture', '')}{company.get('city', '')}"
    cur: Any = data
    for key in dotted.split("."):
        if isinstance(cur, dict):
            cur = cur.get(key)
        else:
            return None
        if cur is None:
            return None
    if isinstance(cur, int):
        return f"{cur}名" if dotted.endswith("employees") else cur
    return cur


def _set_cell_text(cell: _Cell, text: str) -> None:
    """Replace cell text while preserving the cell's paragraph style."""
    if cell.paragraphs:
        first = cell.paragraphs[0]
        for run in list(first.runs):
            run.text = ""
        if first.runs:
            first.runs[0].text = text
        else:
            first.add_run(text)
        # Clear any extra paragraphs in the cell
        for extra in list(cell.paragraphs[1:]):
            extra._element.getparent().remove(extra._element)
    else:
        cell.text = text


# ---------------------------------------------------------------------------
# Append profile-declared tables that aren't already in the template
# ---------------------------------------------------------------------------


def _append_profile_tables(
    doc: _Document, profile: SubsidyProfile, company: dict[str, Any]
) -> list[str]:
    """Append schedule / expense tables at the end of the document.

    A more sophisticated implementation would place each table at the
    section boundary ``place_after_section`` declares; for now, appending
    at the end keeps the official template's structure intact.
    """
    appended: list[str] = []
    for table_spec in profile.tables:
        data = _resolve_path(company, table_spec.data_path) or []
        if not data:
            continue
        _add_table(doc, table_spec, data)
        appended.append(table_spec.table_id)
    return appended


def _add_table(doc: _Document, spec: TableSpec, data: Any) -> None:
    doc.add_paragraph()  # spacer
    heading = doc.add_paragraph()
    heading.add_run(f"【{spec.title}】").bold = True

    rows: list[list[str]] = []
    if spec.table_type == TableType.SCHEDULE:
        for r in data or []:
            rows.append(
                [
                    str(r.get("when", "")),
                    str(r.get("item", "")),
                    str(r.get("content", "")),
                ]
            )
    elif spec.table_type == TableType.EXPENSE_BREAKDOWN:
        for r in data or []:
            rows.append(
                [
                    str(r.get("category", "")),
                    str(r.get("item", "")),
                    f"{r.get('amount', 0):,}",
                ]
            )

    if not rows:
        doc.add_paragraph(f"（{spec.title}：データなし）")
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(spec.columns))
    # Real publisher 様式 docs may not define the "Light Grid Accent 1"
    # built-in style; fall back silently to the document's default style
    # rather than crashing.
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass
    for j, col in enumerate(spec.columns):
        c = table.rows[0].cells[j]
        c.text = col
        for run in c.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val


def _append_profile_charts(
    doc: _Document, profile: SubsidyProfile, company: dict[str, Any]
) -> list[str]:
    """Generate every chart the profile declares and append as figures.

    Charts are rendered to PNG via the existing document_assembler helpers
    so that the official-form path produces visually identical chart art
    to the from-scratch assembler path.
    """
    if not profile.charts:
        return []

    from schemas.subsidy_profile import ChartType
    from tools.document_assembler import _generate_chart  # type: ignore

    appended: list[str] = []
    with tempfile.TemporaryDirectory() as tmp:
        tmp_root = Path(tmp)
        for chart in profile.charts:
            data: Any = None
            cur: Any = company
            for key in chart.data_path.split("."):
                if isinstance(cur, dict):
                    cur = cur.get(key)
                else:
                    cur = None
                    break
            data = cur

            if chart.chart_type == ChartType.EFFECT_BEFORE_AFTER:
                pl = ((company.get("financial") or {}).get("past_3y_pl") or [])
                if pl:
                    expected = (
                        (company.get("planned_project") or {}).get(
                            "expected_outcomes"
                        )
                        or []
                    )
                    before = pl[-1].get("revenue", 0)
                    after = int(before * 1.25)
                    for item in expected:
                        if isinstance(item, dict) and "target_revenue" in item:
                            after = int(item["target_revenue"])
                            break
                    data = {"before": before, "after": after}

            png = tmp_root / f"{chart.chart_id}.png"
            try:
                rendered = _generate_chart(chart, data, png)
            except Exception as e:  # noqa: BLE001
                logger.warning("chart %s render failed: %s", chart.chart_id, e)
                continue
            if not rendered:
                continue
            doc.add_paragraph()
            cap = doc.add_paragraph()
            cap.add_run(f"【{chart.title}】").bold = True
            picture_para = doc.add_paragraph()
            picture_para.add_run().add_picture(
                str(rendered), width=Cm(chart.width_cm)
            )
            appended.append(chart.chart_id)
    return appended


# Convenience for chart insertion when a generated PNG is available
def insert_chart(doc: _Document, png_path: str | Path, width_cm: float = 14.0) -> None:
    para = doc.add_paragraph()
    para.add_run().add_picture(str(png_path), width=Cm(width_cm))
