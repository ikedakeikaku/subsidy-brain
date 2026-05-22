"""Read a 様式 docx and produce a profile-shaped structure.

When the publishing body's actual 様式 .docx is available — either fetched
via GuidelineFetcher (with HTML recovery) or committed by the user to
``templates/<program_id>/`` — its **headings and tables are the most
authoritative source** for the application's section structure. The
publishing body wrote them; nothing else should be trusted over them.

This module walks the docx and emits:

  * ``sections``: every Heading 1/2 found in the document, in order,
    with publisher-specific IDs derived from the heading text.
  * ``applicant_fields``: cells in early tables whose left column
    contains a known applicant label (事業者名 / 代表者氏名 / 住所 / ...).
  * ``tables``: every Word table with column count and first-row headers.

The output is a dict shaped like ``SubsidyProfile`` payloads — caller
hands it to ``ProfileSynthesizer._to_profile()``.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document

logger = logging.getLogger(__name__)


# Heading texts that mark "applicant information" tables — we recognise
# their cells so the profile knows which fields to ask for.
_APPLICANT_LABELS = (
    "事業者名",
    "申請者名",
    "代表者氏名",
    "代表者",
    "事業実施場所",
    "事業所所在地",
    "所在地",
    "従業員数",
    "業種",
    "設立",
    "資本金",
)


_SECTION_NUMBER_RE = re.compile(r"([0-9]+(?:[-\.][0-9]+)*)")

# Real Japanese government docx files often skip the Word Heading style and
# write section titles as plain "Normal" paragraphs with a numbered prefix
# like "０．", "１．", "2-1.", "Ⅰ.". This pattern detects them. Allows
# full-width and half-width digits, optional period, optional sub-number.
_NUMBERED_SECTION_RE = re.compile(
    r"^\s*"
    r"(?:第\s*)?"                           # optional "第" prefix
    r"([0-9０-９IVXⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+(?:[-\.‐－][0-9０-９]+)*)"
    r"[\.\．、:：　 ]+"                      # delimiter
    r"(\S+.*)$"                              # at least one non-space char
)

# Mid-level sub-section markers: （１） （２） ... at the start of the line.
# Used by gov 様式 to group leaves under a top-level numbered section.
_MID_SECTION_RE = re.compile(r"^\s*[（(]\s*([0-9０-９]+)\s*[)）]\s*(\S+.*)$")

# Leaf-level fillable section signal: a heading ending with a char-limit
# marker like "（１０００字以内）" / "（500字程度）" / "（30字以下）".
# When this marker is present, the paragraph IS a fillable input region
# regardless of its style — that's the publishing body's own
# convention across 様式 docs from METI, MoF, MHLW, etc.
_CHAR_LIMIT_RE = re.compile(
    r"[（(]\s*([0-9０-９]+)\s*字\s*(以内|程度|以下|まで)\s*[）)]"
)

# Circled-number / circled-text prefix for leaves where the publisher used
# auto-numbered lists. python-docx's .text hides the auto-number, so we
# also match explicit ①②③ and "③－２"-style explicit numbering.
_CIRCLED_NUM_RE = re.compile(r"^\s*[①-⑳㉑-㉟]")
_EXPLICIT_SUBNUM_RE = re.compile(r"^\s*[①-⑳][－\-‐]\s*[0-9０-９]+")


def read_form_docx(docx_path: str | Path) -> dict[str, Any] | None:
    """Parse a 様式 docx and return a profile-shaped dict.

    Returns ``None`` if the file is missing or not a valid docx package.
    """
    p = Path(docx_path)
    if not p.exists() or p.stat().st_size < 1024:
        return None

    # Validate it's actually a docx — caller may have passed something
    # that GuidelineFetcher saved with a .docx name but is actually HTML.
    if not _is_real_docx(p):
        return None

    try:
        doc = Document(str(p))
    except Exception as e:  # noqa: BLE001
        logger.warning("docx_structure_reader: open failed for %s: %s", p, e)
        return None

    sections: list[dict[str, Any]] = []
    applicant_fields: list[dict[str, str]] = []
    tables_meta: list[dict[str, Any]] = []

    seen_ids: set[str] = set()
    # Track the current top-level so leaves can disambiguate their IDs
    # when their visible text is duplicated across siblings.
    current_top: str | None = None
    leaf_index_in_top: dict[str, int] = {}

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = (para.text or "").strip()
        if not text:
            continue

        classification = _classify_heading(para, style, text)
        if classification is None:
            continue
        kind, char_limit = classification

        section_id = _derive_section_id(text)
        # For auto-numbered leaves (List Paragraph style, ①②③ hidden by
        # python-docx), the visible text may not contain enough information
        # to disambiguate (e.g. multiple "...（１０００字以内）" lines).
        # Append an index keyed off the containing top-level section so
        # ids stay stable but unique across the document.
        if kind == "leaf" and current_top:
            idx = leaf_index_in_top.get(current_top, 0) + 1
            leaf_index_in_top[current_top] = idx
            if section_id in seen_ids:
                section_id = f"{current_top}_leaf_{idx:02d}"
        if section_id in seen_ids:
            section_id = f"{section_id}_{len(seen_ids)}"
        seen_ids.add(section_id)

        if kind == "leaf":
            # Char limit drives target/min/max precisely (no heuristic
            # guesswork — the publisher told us the limit).
            limit = char_limit or _guess_target_chars(text)
            target = int(limit * 0.9)
            sections.append(
                {
                    "section_id": section_id,
                    "display_name": text,
                    "target_chars": target,
                    "min_chars": int(limit * 0.5),
                    "max_chars": limit,
                    "kind": "leaf",
                }
            )
        elif kind == "container":
            sections.append(
                {
                    "section_id": section_id,
                    "display_name": text,
                    "target_chars": 0,
                    "min_chars": 0,
                    "max_chars": 0,
                    "kind": "container",
                }
            )
        else:  # "section" (top-level)
            target = _guess_target_chars(text)
            sections.append(
                {
                    "section_id": section_id,
                    "display_name": text,
                    "target_chars": target,
                    "min_chars": int(target * 0.65),
                    "max_chars": int(target * 1.3),
                    "kind": "section",
                }
            )
            current_top = section_id
            leaf_index_in_top[section_id] = 0

    for table in doc.tables:
        cols = len(table.columns)
        rows = len(table.rows)
        if not rows or not cols:
            continue
        header_row = table.rows[0]
        header_cells = [c.text.strip() for c in header_row.cells]

        # Detect applicant-info tables: cell in column 0 contains a known label
        for row in table.rows:
            label = row.cells[0].text.strip() if row.cells else ""
            if any(lbl in label for lbl in _APPLICANT_LABELS):
                applicant_fields.append(
                    {"label": label, "field_id": _label_to_field_id(label)}
                )

        tables_meta.append(
            {
                "rows": rows,
                "cols": cols,
                "first_row": header_cells,
            }
        )

    if not sections:
        # No headings found — empty form? Treat as failure so caller falls
        # through to PDF / web_search.
        logger.info("docx_structure_reader: no headings in %s", p)
        return None

    return {
        "source": "form_docx",
        "form_path": str(p),
        "sections": sections,
        "applicant_fields": applicant_fields,
        "tables": tables_meta,
    }


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _is_real_docx(path: Path) -> bool:
    import zipfile

    try:
        with zipfile.ZipFile(path) as z:
            return "word/document.xml" in z.namelist()
    except (zipfile.BadZipFile, OSError):
        return False


def _classify_heading(
    para: Any, style: str, text: str
) -> tuple[str, int | None] | None:
    """Return ``(kind, char_limit)`` if the paragraph is a heading, else None.

    ``kind`` is one of:
      * ``"leaf"`` — fillable input region with a char-count marker. Always
        rendered with its own body, sized to the publisher's char limit.
      * ``"container"`` — mid-level grouping like "（１）申請者の概要".
        No body of its own; its leaves carry the content.
      * ``"section"`` — top-level numbered section ("０．", "１．", "2-1.").
        Has body when no leaves are nested under it (e.g. "０．誓約・同意事項").
    """
    if not text or len(text) > 200:
        # Heading text is bounded — anything longer is body prose.
        return None

    char_limit = _extract_char_limit(text)

    # Strongest signal: a char-limit marker — this paragraph is a leaf
    # regardless of style or numbering prefix. The publisher chose to put
    # an "（NNN字以内）" annotation here precisely because it's a fillable
    # input region. Style filter is intentionally loose so we catch the
    # auto-numbered List Paragraph leaves whose visible ①②③ are hidden
    # by python-docx's run-text view.
    if char_limit is not None:
        return ("leaf", char_limit)

    # Explicit circled-number or "③－２" prefix → leaf without char limit.
    if _CIRCLED_NUM_RE.match(text) or _EXPLICIT_SUBNUM_RE.match(text):
        # Still need a sanity bound — body prose can start with ① too.
        if len(text) <= 80 and "。" not in text[:60]:
            return ("leaf", None)

    # Mid-level container: "（１）", "（２）", ...
    # Parenthetical notes inside the heading often contain "。" (e.g.
    # "（３）役員一覧（監査役を含む。）"), so the punctuation gate is only
    # applied outside parentheses.
    if _MID_SECTION_RE.match(text):
        outside_parens = re.sub(r"[（(][^（）()]*[）)]", "", text)
        if "、" not in outside_parens and "。" not in outside_parens:
            return ("container", None)

    # Word Heading / 見出し style — treat as top-level section.
    if style.startswith("Heading") or style.startswith("見出し"):
        return ("section", None)

    # Numbered top-level: "０．", "１．", "2-1.", "Ⅰ."
    if len(text) > 80:
        return None
    if not _NUMBERED_SECTION_RE.match(text):
        return None
    head = text[:60]
    if "、" in head or "。" in head:
        return None
    return ("section", None)


def _extract_char_limit(text: str) -> int | None:
    """Return the integer char limit from a heading like "...（１０００字以内）",
    or ``None`` if no char-limit marker is present."""
    m = _CHAR_LIMIT_RE.search(text)
    if not m:
        return None
    raw = m.group(1).translate(_FULLWIDTH_DIGITS)
    try:
        return int(raw)
    except ValueError:
        return None


_FULLWIDTH_DIGITS = str.maketrans("０１２３４５６７８９", "0123456789")
_ROMAN_TO_INT = {
    "Ⅰ": "1", "Ⅱ": "2", "Ⅲ": "3", "Ⅳ": "4", "Ⅴ": "5",
    "Ⅵ": "6", "Ⅶ": "7", "Ⅷ": "8", "Ⅸ": "9", "Ⅹ": "10",
}


def _normalise_number(token: str) -> str:
    """Convert full-width or Roman section numbers to ASCII digits."""
    token = token.translate(_FULLWIDTH_DIGITS)
    for roman, ascii_n in _ROMAN_TO_INT.items():
        token = token.replace(roman, ascii_n)
    return token


def _derive_section_id(heading_text: str) -> str:
    """Stable ASCII id for a heading. Falls back to a slugified hash of
    the text if no section number is present."""
    # Look at the start of the heading first (gov-style numbered sections)
    nm = _NUMBERED_SECTION_RE.match(heading_text)
    if nm:
        num = _normalise_number(nm.group(1)).replace(".", "_").replace("-", "_")
        return f"section_{num}"
    match = _SECTION_NUMBER_RE.search(heading_text)
    if match:
        num = match.group(1).replace(".", "_").replace("-", "_")
        return f"section_{num}"
    # Tag bonus / 加点 headings
    if "加点" in heading_text or "重点政策" in heading_text:
        # Try to extract the kind from the heading
        for keyword, suffix in (
            ("事業環境変化", "env_change"),
            ("賃金", "wage_increase"),
            ("賃上げ", "wage_increase"),
            ("赤字", "deficit"),
            ("卒業", "graduation"),
            ("創業", "startup"),
            ("再生", "rehabilitation"),
            ("地域", "regional"),
        ):
            if keyword in heading_text:
                return f"bonus_{suffix}"
        return "bonus_other"
    # Slugify any other heading
    slug = re.sub(r"\W+", "_", heading_text.lower())[:32].strip("_")
    return slug or "section_unknown"


def _guess_target_chars(heading_text: str) -> int:
    """Heuristic per-section target. Larger sections (経営計画 / 補助事業
    の取組内容 等) get higher targets; small ones (経費明細 / 事業名)
    smaller."""
    if "今後のプラン" in heading_text or "取組内容" in heading_text:
        return 1500
    if "経営" in heading_text or "計画" in heading_text:
        return 900
    if "効果" in heading_text or "市場" in heading_text:
        return 800
    if "強み" in heading_text or "弱み" in heading_text:
        return 700
    if "概要" in heading_text:
        return 700
    if "加点" in heading_text:
        return 500
    if "事業名" in heading_text:
        return 100
    if "経費" in heading_text:
        return 400
    return 600


def _label_to_field_id(label: str) -> str:
    mapping = {
        "事業者名": "company_name",
        "申請者名": "company_name",
        "代表者氏名": "representative",
        "代表者": "representative",
        "事業実施場所": "business_address",
        "事業所所在地": "business_address",
        "所在地": "business_address",
        "従業員数": "employee_count",
        "業種": "industry",
        "設立": "founded",
        "資本金": "capital",
    }
    for k, v in mapping.items():
        if k in label:
            return v
    return re.sub(r"\W+", "_", label.lower())
