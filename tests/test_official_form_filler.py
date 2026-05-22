"""Tests for OfficialFormFiller.

Real 様式 docx files distributed by Japanese subsidy bodies don't carry
``{{placeholder}}`` markers — they're shipped as form-style documents
with headings and empty body space. The filler has to recognise the
structure on its own. These tests build a doc that matches that shape and
verify every fill axis: section bodies, applicant info cells, appended
profile tables, and appended profile charts.
"""
from __future__ import annotations

import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


# ---------------------------------------------------------------------------
# Helpers — build an "official-style" template matching the shape of a real 様式2
# ---------------------------------------------------------------------------


def _build_official_style_template(out: Path) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)
    doc.sections[0].left_margin = Cm(2.5)

    title = doc.add_heading("様式２（経営計画書 兼 補助事業計画書）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Applicant info table (labels in col 0, empty value cells in col 1)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, label in enumerate(["事業者名", "代表者氏名", "事業実施場所", "従業員数"]):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = ""  # empty value cell

    # Bonus block
    doc.add_heading("＜加点項目＞", level=1)
    doc.add_heading("【重点政策加点】事業環境変化加点", level=2)
    doc.add_paragraph()
    doc.add_paragraph()

    # Body sections
    for heading in [
        "1-1. 自社の概要",
        "1-2. 売上・利益の状況",
        "1-3. 経営課題",
        "2-1. 市場の動向",
        "3. 強み・弱み",
        "4-2. 今後のプラン",
        "補助事業の効果",
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph()

    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_official_form_filler_fills_every_section(tmp_path: Path) -> None:
    from demo.mock_story import MOCK_STORY
    from schemas.subsidy_profile import load_profile
    from tools.official_form_filler import fill_official_form

    template = _build_official_style_template(tmp_path / "様式2.docx")
    out = tmp_path / "filled.docx"

    profile = load_profile(ROOT / "demo" / "sample_profile.yaml")
    company = yaml.safe_load(
        (ROOT / "demo" / "sample_company.yaml").read_text(encoding="utf-8")
    )

    report = fill_official_form(
        template_path=template,
        out_path=out,
        profile=profile,
        story=MOCK_STORY,
        company=company,
    )

    # Every section in the profile should have been located in the template.
    assert report["sections_not_found"] == [], report
    assert set(report["sections_filled"]) == {s.section_id for s in profile.sections}

    # All four applicant-info cells should have been filled.
    assert report["applicant_cells_filled"] == 4

    # Profile tables and charts should have been appended.
    assert "table_schedule" in report["tables_appended"]
    assert "table_expense" in report["tables_appended"]
    assert "chart_revenue_trend" in report.get("charts_appended", [])

    # The output is a real docx, larger than the template (because we
    # appended figures and tables).
    assert out.exists()
    assert out.stat().st_size > template.stat().st_size


def test_official_form_filler_preserves_template_text(tmp_path: Path) -> None:
    """Filling must not erase the template's existing headings or labels."""
    from demo.mock_story import MOCK_STORY
    from schemas.subsidy_profile import load_profile
    from tools.official_form_filler import fill_official_form

    template = _build_official_style_template(tmp_path / "様式2.docx")
    out = tmp_path / "filled.docx"

    profile = load_profile(ROOT / "demo" / "sample_profile.yaml")
    company = yaml.safe_load(
        (ROOT / "demo" / "sample_company.yaml").read_text(encoding="utf-8")
    )

    fill_official_form(
        template_path=template,
        out_path=out,
        profile=profile,
        story=MOCK_STORY,
        company=company,
    )

    filled = Document(str(out))
    text_blob = "\n".join(p.text for p in filled.paragraphs)
    # Original template headings must survive
    assert "様式２" in text_blob
    assert "1-1. 自社の概要" in text_blob
    assert "4-2. 今後のプラン" in text_blob
    # Original applicant labels must survive in cells
    labels_found = set()
    for table in filled.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip() in {"事業者名", "代表者氏名", "事業実施場所", "従業員数"}:
                    labels_found.add(cell.text.strip())
    assert labels_found >= {"事業者名", "代表者氏名"}


def test_section_signature_matches_variant_headings() -> None:
    """The signature should treat '4-2. 今後のプラン' and
    '4-2. 今後のプラン（施策）' as the same section."""
    from tools.official_form_filler import _section_signature

    a = _section_signature("4-2. 今後のプラン")
    b = _section_signature("4-2. 今後のプラン（施策）")
    assert a[0] == b[0] == "4-2"

    # And it should recognise bonus headings even without a number
    bonus_a = _section_signature("【重点政策加点】事業環境変化加点")
    bonus_b = _section_signature("【加点】事業環境変化")
    assert bonus_a[0] == bonus_b[0] == "bonus:env_change"
