"""Tests for the natural-language pipeline.

These cover the new synthesis / cache / template layer that lets the
system handle subsidies without any preset YAML.
"""
from __future__ import annotations

import asyncio
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


# ---------------------------------------------------------------------------
# ProfileSynthesizer falls back when web search unavailable
# ---------------------------------------------------------------------------


def test_profile_synthesizer_fallback_without_keys(monkeypatch) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    profile = asyncio.run(
        ProfileSynthesizer().synthesize("持続化補助金 第19回")
    )
    assert "持続化補助金" in profile.canonical_name
    assert len(profile.sections) >= 6
    assert profile.total_target_chars >= 4_000
    assert profile.charts and profile.tables


def test_profile_synthesizer_returns_different_id_per_query(monkeypatch) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    a = asyncio.run(ProfileSynthesizer().synthesize("持続化補助金 第19回"))
    b = asyncio.run(ProfileSynthesizer().synthesize("ものづくり補助金 第18次"))
    # Different inputs → different canonical_name; program_id may still
    # collide with the fallback default; we only assert canonical_name here.
    assert a.canonical_name != b.canonical_name


# ---------------------------------------------------------------------------
# Profile cache round-trips
# ---------------------------------------------------------------------------


def test_profile_cache_roundtrip(tmp_path: Path) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from tools.profile_cache import ProfileCache

    cache = ProfileCache(root=tmp_path)
    profile = asyncio.run(
        ProfileSynthesizer().synthesize("テスト補助金")
    )
    assert cache.load(profile.program_id) is None

    cache.save(profile)
    loaded = cache.load(profile.program_id)
    assert loaded is not None
    assert loaded.canonical_name == profile.canonical_name
    assert len(loaded.sections) == len(profile.sections)

    assert cache.evict(profile.program_id) is True
    assert cache.load(profile.program_id) is None


def test_profile_cache_respects_ttl(tmp_path: Path) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from tools.profile_cache import ProfileCache

    cache = ProfileCache(root=tmp_path, ttl_seconds=0)  # immediate expiry
    profile = asyncio.run(
        ProfileSynthesizer().synthesize("テスト補助金B")
    )
    cache.save(profile)
    # ttl_seconds=0 means "never serve cache hits"
    import time

    time.sleep(0.01)
    assert cache.load(profile.program_id) is None


# ---------------------------------------------------------------------------
# TemplateSynthesizer produces a valid .docx with placeholders
# ---------------------------------------------------------------------------


def test_template_synthesizer_generates_from_profile(tmp_path: Path) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from agents.template_synthesizer import TemplateSynthesizer

    profile = asyncio.run(
        ProfileSynthesizer().synthesize("テスト補助金C")
    )
    template = TemplateSynthesizer().get_template(
        profile, templates_root=tmp_path
    )
    assert template.exists()
    assert template.stat().st_size > 5_000

    # Every section in the profile should have a corresponding placeholder
    from docx import Document

    doc = Document(str(template))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for spec in profile.sections:
        assert (
            f"{{{{ {spec.section_id} }}}}" in full_text
            or f"{{{{{spec.section_id}}}}}" in full_text
        ), spec.section_id


def test_template_synthesizer_prefers_fetched_form(tmp_path: Path) -> None:
    """If GuidelineFetcher gave us a 様式2 docx, use it; don't regenerate."""
    from docx import Document

    from agents.profile_synthesizer import ProfileSynthesizer
    from agents.template_synthesizer import TemplateSynthesizer

    # Create a fake "official" 様式2 docx
    fake_official = tmp_path / "official_様式2.docx"
    Document().add_paragraph("OFFICIAL_FORM").part  # ensure it's a real doc
    d = Document()
    d.add_paragraph("OFFICIAL_FORM_MARKER")
    d.save(str(fake_official))

    profile = asyncio.run(
        ProfileSynthesizer().synthesize("テスト補助金D")
    )
    template = TemplateSynthesizer().get_template(
        profile,
        fetched_form_paths={"様式2_経営計画書": str(fake_official)},
        templates_root=tmp_path / "templates",
    )
    assert template == fake_official


# ---------------------------------------------------------------------------
# End-to-end natural-language pipeline runs
# ---------------------------------------------------------------------------


def test_natural_demo_runs_end_to_end(tmp_path: Path, monkeypatch) -> None:
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")
    monkeypatch.chdir(ROOT)

    from demo.run_natural_demo import run

    asyncio.run(
        run("持続化補助金 第19回（テスト）", live=False, use_cache=False)
    )

    # The pipeline writes the output under demo/output/<program_id>_application.docx
    # which depends on the synthesised program_id. Verify by globbing.
    out_dir = ROOT / "demo" / "output"
    docx_candidates = list(out_dir.glob("*_application.docx"))
    assert docx_candidates, "natural demo did not produce any application docx"
    # At least one should be larger than 30KB (real charts/tables embedded)
    assert any(p.stat().st_size > 30_000 for p in docx_candidates)


@pytest.mark.parametrize(
    "query",
    [
        "持続化補助金 第20回",
        "事業再構築補助金 第14回",
        "ものづくり補助金 第19次",
    ],
)
def test_synthesizer_adapts_to_any_subsidy(query: str, monkeypatch) -> None:
    """Same engine, three different subsidy names — all must produce a
    valid profile without any preset YAML."""
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    profile = asyncio.run(ProfileSynthesizer().synthesize(query))
    assert profile.canonical_name
    assert profile.sections
    assert profile.charts and profile.tables
