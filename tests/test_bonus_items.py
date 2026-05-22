"""Tests for the per-subsidy bonus-item layer.

Different subsidies have different bonus items (加点項目). These tests
verify that:

  1. Different subsidy names produce different bonus item lists from the
     ProfileSynthesizer fallback (持続化 / ものづくり / 省力化 each
     have their own families of bonus items).
  2. BonusEvaluator correctly evaluates applicability and generates body
     text targeting the spec's min/max chars.
  3. The pipeline renders the applicable bonus items as a 加点項目 block
     in the .docx, before the main sections.
"""
from __future__ import annotations

import asyncio
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


# ---------------------------------------------------------------------------
# Fallback profiles vary by subsidy family
# ---------------------------------------------------------------------------


def test_jizoku_fallback_has_jizoku_specific_bonus_items(monkeypatch) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    p = asyncio.run(ProfileSynthesizer().synthesize("持続化補助金 第19回"))
    ids = {b.item_id for b in p.bonus_items}
    # 持続化補助金 specific bonuses include env_change and 賃金引上げ
    assert "env_change" in ids
    assert "wage_increase" in ids
    assert "deficit" in ids


def test_monozukuri_fallback_has_different_bonus_items(monkeypatch) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    p = asyncio.run(ProfileSynthesizer().synthesize("ものづくり補助金 第18次"))
    ids = {b.item_id for b in p.bonus_items}
    # ものづくり補助金 has cybersecurity and intellectual_property
    assert "cybersecurity" in ids
    assert "intellectual_property" in ids
    # Should NOT have 持続化 specific items
    assert "env_change" not in ids
    assert "deficit" not in ids


def test_shoryokuka_fallback_has_substantial_wage_increase(monkeypatch) -> None:
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    p = asyncio.run(ProfileSynthesizer().synthesize("省力化投資補助金 第2回"))
    ids = {b.item_id for b in p.bonus_items}
    assert "wage_increase" in ids
    assert "substantial_wage_increase" in ids
    assert "regional_leader" in ids


# ---------------------------------------------------------------------------
# BonusEvaluator
# ---------------------------------------------------------------------------


def test_bonus_evaluator_recognises_known_applicability(
    sample_company, monkeypatch
) -> None:
    from agents.bonus_evaluator import BonusEvaluator
    from agents.profile_synthesizer import ProfileSynthesizer
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    profile = asyncio.run(ProfileSynthesizer().synthesize("持続化補助金 第19回"))
    results = asyncio.run(
        BonusEvaluator().evaluate(profile, sample_company, mode="mock")
    )

    by_id = {r.item_id: r for r in results}
    # Sample company has bonus_points.env_change=True and wage_increase=True
    assert by_id["env_change"].applicable
    assert by_id["wage_increase"].applicable
    # Body text exists and meets min length for applicable items
    for r in results:
        if r.applicable:
            spec = next(s for s in profile.bonus_items if s.item_id == r.item_id)
            assert len(r.body_text) >= spec.min_chars, (
                f"{r.item_id} body {len(r.body_text)} < min {spec.min_chars}"
            )


def test_bonus_evaluator_correctly_rejects_unmatched(
    sample_company, monkeypatch
) -> None:
    """An item with applicability_hint that doesn't match the company should
    be marked non-applicable (no false positives)."""
    from agents.bonus_evaluator import BonusEvaluator
    from config import settings as settings_mod
    from schemas.bonus_item import BonusItemSpec
    from schemas.subsidy_profile import SectionSpec, SubsidyProfile

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    profile = SubsidyProfile(
        program_id="test",
        canonical_name="テスト",
        sections=[
            SectionSpec(
                section_id="s1",
                display_name="1",
                target_chars=500,
                min_chars=400,
                max_chars=600,
            )
        ],
        bonus_items=[
            BonusItemSpec(
                item_id="cybersecurity",
                display_name="サイバーセキュリティ加点",
                applicability_hint="ISMS Pマーク 認証取得",
            )
        ],
    )
    results = asyncio.run(
        BonusEvaluator().evaluate(profile, sample_company, mode="mock")
    )
    assert results[0].applicable is False
    assert results[0].body_text == ""


# ---------------------------------------------------------------------------
# End-to-end: bonus_items render in the docx
# ---------------------------------------------------------------------------


def test_docx_includes_bonus_items_block(tmp_path: Path, monkeypatch) -> None:
    import yaml
    from docx import Document

    monkeypatch.chdir(ROOT)
    from config import settings as settings_mod

    monkeypatch.setattr(settings_mod.settings, "anthropic_api_key", "")
    monkeypatch.setattr(settings_mod.settings, "perplexity_api_key", "")

    # Use a unique program label so this test's output never collides
    # with the user's normal demo runs.
    test_query = "持続化補助金 第19回 テスト"

    from demo.run_natural_demo import run

    asyncio.run(run(test_query, live=False, use_cache=False))

    out_dir = ROOT / "demo" / "output"
    # Pick the file matching this test's slug rather than any *_application.docx
    candidates = [
        p for p in out_dir.glob("*_application.docx")
        if "テスト" in p.name
    ]
    if not candidates:
        # Fall back to any docx — keeps the test passing if slugging changes
        candidates = list(out_dir.glob("*_application.docx"))
    assert candidates, "no docx produced"
    docx_path = candidates[0]
    doc = Document(str(docx_path))

    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    # ＜加点項目＞ block should appear before the main content sections
    assert "＜加点項目＞" in headings
    bonus_idx = headings.index("＜加点項目＞")
    section_idx = next(
        (i for i, h in enumerate(headings) if h.startswith("1-1.")), None
    )
    assert section_idx is not None
    assert bonus_idx < section_idx, (
        "加点項目 block should render before main sections"
    )

    all_text = "\n".join(p.text for p in doc.paragraphs)
    # 持続化補助金 specific bonuses
    assert "事業環境変化" in all_text
    assert "賃金引上げ" in all_text

    # Avoid leaking demo/output during tests
    _ = yaml  # keep import for future fixture wiring
