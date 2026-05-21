"""Generate the fictional sample template .docx used by the demo.

We don't ship a binary .docx in version control; we generate it on first run.
This keeps the repository diff-friendly and ensures the template's structure
is auditable as Python code rather than as an opaque binary.

Running this script (which the demo does automatically) produces
``templates/sample_hanro_kaitaku_v1/様式2.docx`` containing placeholder
tokens that ``tools.template_filler.fill_template`` can substitute.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent / "sample_hanro_kaitaku_v1"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def build_form_2() -> Path:
    """Build a deliberately styled mock of an official 様式2 (経営計画書).

    Real official forms have heavy table layouts, fixed margins, and a
    cover-page header block. This mock reproduces those traits at a level
    high enough for the demo to demonstrate template fidelity preservation.
    """
    out_path = OUT_DIR / "様式2.docx"
    doc = Document()

    # Page setup (A4, 25mm margins — close to the actual official forms)
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)

    # Title block
    title = doc.add_heading("様式２（経営計画書 兼 補助事業計画書）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.add_run("販路開拓支援補助金（架空）／第1回").bold = True

    # Applicant header table — preserves a 2-column layout typical of 様式
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Light Grid Accent 1"
    hdr.rows[0].cells[0].text = "事業者名"
    hdr.rows[0].cells[1].text = "{{ company_name }}"
    hdr.rows[1].cells[0].text = "代表者氏名"
    hdr.rows[1].cells[1].text = "{{ representative }}"
    hdr.rows[2].cells[0].text = "事業実施場所"
    hdr.rows[2].cells[1].text = "{{ business_address }}"
    hdr.rows[3].cells[0].text = "従業員数"
    hdr.rows[3].cells[1].text = "{{ employee_count }}"

    # ＜加点項目＞
    doc.add_heading("＜加点項目＞", level=1)
    doc.add_heading("【重点政策加点】事業環境変化加点", level=2)
    doc.add_paragraph("{{ bonus_env_change }}")

    # ＜経営計画＞
    doc.add_heading("＜経営計画＞", level=1)

    doc.add_heading("1-1. 自社の概要", level=2)
    doc.add_paragraph("{{ section_1_1 }}")

    doc.add_heading("1-2. 売上・利益の状況", level=2)
    doc.add_paragraph("{{ section_1_2 }}")
    # 売上推移テーブル
    pl_tbl = doc.add_table(rows=4, cols=3)
    pl_tbl.style = "Light Grid Accent 1"
    pl_tbl.rows[0].cells[0].text = "年度"
    pl_tbl.rows[0].cells[1].text = "売上（円）"
    pl_tbl.rows[0].cells[2].text = "営業利益（円）"
    for i, year_key in enumerate(["pl_y1", "pl_y2", "pl_y3"], start=1):
        pl_tbl.rows[i].cells[0].text = f"{{{{ {year_key}_year }}}}"
        pl_tbl.rows[i].cells[1].text = f"{{{{ {year_key}_revenue }}}}"
        pl_tbl.rows[i].cells[2].text = f"{{{{ {year_key}_profit }}}}"

    doc.add_heading("1-3. 経営課題", level=2)
    doc.add_paragraph("{{ section_1_3 }}")

    # ＜補助事業計画＞
    doc.add_heading("＜補助事業計画＞", level=1)
    doc.add_heading("4-2. 今後のプラン", level=2)
    doc.add_paragraph("{{ section_4_2 }}")

    doc.add_heading("補助事業の効果", level=2)
    doc.add_paragraph("{{ section_effect }}")

    # ＜経費明細＞
    doc.add_heading("＜経費明細＞", level=1)
    exp_tbl = doc.add_table(rows=2, cols=3)
    exp_tbl.style = "Light Grid Accent 1"
    exp_tbl.rows[0].cells[0].text = "区分"
    exp_tbl.rows[0].cells[1].text = "明細"
    exp_tbl.rows[0].cells[2].text = "金額（円）"
    exp_tbl.rows[1].cells[0].text = "{{ expense_table_csv }}"
    exp_tbl.rows[1].cells[1].text = ""
    exp_tbl.rows[1].cells[2].text = ""

    funding = doc.add_paragraph()
    funding.add_run("補助金申請額: {{ subsidy_amount }} 円\n").bold = True
    funding.add_run("自己負担額  : {{ self_funding }} 円")

    doc.save(str(out_path))
    return out_path


def main() -> None:
    p = build_form_2()
    print(f"✓ Built sample template: {p}")


if __name__ == "__main__":
    main()
