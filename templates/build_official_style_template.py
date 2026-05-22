"""Generate a docx that mimics how a real subsidy 様式2 is structured.

Real official forms distributed by 全国商工会連合会 / 中小企業庁 / etc.
share a common shape:

  * Cover page with title and the 様式 number.
  * Applicant information table at the top: labels in the left column
    ("事業者名" / "代表者氏名" / etc.), empty cells in the right column.
  * Section headings ("1-1. 自社の概要" / "4-2. 今後のプラン" / etc.)
    each followed by a blank space (sometimes several blank paragraphs)
    where the applicant writes the body.
  * Pre-formatted tables for schedules and expense breakdowns appended
    at the end.

Critically: **no placeholder markers**. The applicant just clicks into
the empty paragraph after each heading and types. This script generates a
file with that exact shape so the OfficialFormFiller can be demoed end-
to-end on something that resembles the real thing.

In production, this script is NOT used — the official .docx downloaded
by ``GuidelineFetcher`` takes its place.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent / "official_style_sample"
OUT_DIR.mkdir(parents=True, exist_ok=True)


def build_official_style_form() -> Path:
    out_path = OUT_DIR / "様式2_経営計画書.docx"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Hiragino Sans"
    style.font.size = Pt(10.5)

    # ── Cover ────────────────────────────────────────────────────
    title = doc.add_heading("様式２（経営計画書 兼 補助事業計画書）", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sub.add_run("補助金番号: ＿＿＿＿＿＿＿＿＿＿＿＿").italic = True

    # ── Applicant info ───────────────────────────────────────────
    doc.add_paragraph()
    hdr = doc.add_table(rows=4, cols=2)
    hdr.style = "Light Grid Accent 1"
    for i, label in enumerate(["事業者名", "代表者氏名", "事業実施場所", "従業員数"]):
        hdr.rows[i].cells[0].text = label
        hdr.rows[i].cells[1].text = ""  # empty value cell

    # ── Bonus block ──────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("＜加点項目＞", level=1)
    doc.add_heading("【重点政策加点】事業環境変化加点", level=2)
    doc.add_paragraph()  # empty body
    doc.add_paragraph()  # extra blank

    # ── 経営計画 ─────────────────────────────────────────────────
    doc.add_heading("＜経営計画＞", level=1)
    for heading in [
        "1-1. 自社の概要",
        "1-2. 売上・利益の状況",
        "1-3. 経営課題",
        "2-1. 市場の動向",
        "3. 強み・弱み",
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph()
        doc.add_paragraph()

    # ── 補助事業計画 ─────────────────────────────────────────────
    doc.add_heading("＜補助事業計画＞", level=1)
    for heading in ["4-2. 今後のプラン", "補助事業の効果"]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph()
        doc.add_paragraph()

    doc.save(str(out_path))
    return out_path


def main() -> None:
    p = build_official_style_form()
    print(f"✓ Built official-style sample: {p} ({p.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
